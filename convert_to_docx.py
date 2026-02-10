"""
Convert USER_GUIDE.md to USER_GUIDE.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_user_guide():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Attendance Dashboard - User Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('For HR & Payroll Staff')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True

    doc.add_paragraph()

    # ========== GETTING STARTED ==========
    doc.add_heading('Getting Started', level=1)

    doc.add_heading('Step 1: Open the Application', level=2)
    doc.add_paragraph('Double-click on AttendanceDashboard.exe to start the application.')
    doc.add_paragraph('A window will appear with the title "Attendance Dashboard v2.3"')

    # ========== LOADING FILES ==========
    doc.add_heading('Loading Your Files', level=1)
    doc.add_paragraph('You need to load 3 files before generating a report:')

    # Step 2
    doc.add_heading('Step 2: Load Attendance Files', level=2)
    doc.add_paragraph('1. Click the "Select Attendance Files" button')
    doc.add_paragraph('2. Navigate to your attendance files folder')
    doc.add_paragraph('3. Select one or more .xlsx files (hold Ctrl to select multiple)')
    doc.add_paragraph('4. Click "Open"')

    p = doc.add_paragraph()
    p.add_run('What are attendance files? ').bold = True
    p.add_run('Excel files exported from your attendance machine containing clock in/out times.')

    # Step 3
    doc.add_heading('Step 3: Load Master Data', level=2)
    doc.add_paragraph('1. Click the "Select Master Data" button')
    doc.add_paragraph('2. Select your employee master data Excel file')
    doc.add_paragraph('3. Click "Open"')

    p = doc.add_paragraph()
    p.add_run('What is master data? ').bold = True
    p.add_run('Excel file with employee information: CRM, Name, Department, National ID, Vendor, PS ID, Join Date.')

    # Step 4
    doc.add_heading('Step 4: Load Leave Sheet (Optional but Recommended)', level=2)
    doc.add_paragraph('1. Click the "Select Leave Sheet" button')
    doc.add_paragraph('2. Select your leave records Excel file')
    doc.add_paragraph('3. Click "Open"')

    p = doc.add_paragraph()
    p.add_run('What is the leave sheet? ').bold = True
    p.add_run('Excel file with monthly sheets (Jan, Feb, etc.) showing who took Annual Leave, Sick Leave, etc.')

    # ========== FILTERS ==========
    doc.add_heading('Using Employee Filters', level=1)
    doc.add_paragraph('After loading Master Data, you can filter which employees appear in the report.')

    doc.add_heading('Filter by Department', level=2)
    doc.add_paragraph('Check the box to include that department, uncheck to exclude.')
    doc.add_paragraph('Use "Select All" to quickly select/deselect all departments.')

    doc.add_heading('Filter by CRM', level=2)
    doc.add_paragraph('Same as department filter - check to include, uncheck to exclude.')

    doc.add_heading('Reset Filters', level=2)
    doc.add_paragraph('Click the "Reset" button to restore all selections.')

    # ========== GENERATING REPORT ==========
    doc.add_heading('Generating the Report', level=1)

    doc.add_heading('Step 5: Generate Report', level=2)
    doc.add_paragraph('1. Make sure all required files are loaded (green checkmarks)')
    doc.add_paragraph('2. Click the "Generate Report" button')
    doc.add_paragraph('3. Choose where to save the Excel file')
    doc.add_paragraph('4. Wait for processing to complete')

    # ========== UNDERSTANDING REPORT ==========
    doc.add_heading('Understanding the Generated Excel Report', level=1)
    doc.add_paragraph('The report contains 4 sheets:')

    # Sheet 1
    doc.add_heading('Sheet 1: Summary Report', level=2)
    doc.add_paragraph('This is your main working sheet. Shows employee attendance with dates as columns.')

    # Color Guide Table
    doc.add_heading('Color Guide', level=3)
    color_table = doc.add_table(rows=9, cols=2)
    color_table.style = 'Table Grid'

    colors = [
        ('Color', 'Meaning'),
        ('Green', 'No penalty (Normal, Approved, Leaves)'),
        ('Yellow', 'Late (with penalty)'),
        ('Red', 'Absent (2 days deduction)'),
        ('Pink', 'Missing Punch'),
        ('Orange', 'Half Day / Early Departure'),
        ('Light Blue', 'Sick Leave / Unpaid Leave'),
        ('Purple', 'Backdated Leave (from previous month)'),
        ('Gray', 'Weekend / OFF day'),
    ]

    color_codes = {
        'Green': '90EE90',
        'Yellow': 'FFFF00',
        'Red': 'FF6B6B',
        'Pink': 'FFB6C1',
        'Orange': 'FFA500',
        'Light Blue': 'ADD8E6',
        'Purple': 'E1D5E7',
        'Gray': 'D3D3D3',
    }

    for i, (color, meaning) in enumerate(colors):
        row = color_table.rows[i]
        row.cells[0].text = color
        row.cells[1].text = meaning
        if i == 0:
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True
        elif color in color_codes:
            set_cell_shading(row.cells[0], color_codes[color])

    doc.add_paragraph()

    # Sheet 2
    doc.add_heading('Sheet 2: Individual Analytics', level=2)
    doc.add_paragraph('Shows attendance statistics for each employee (Normal count, Late count, Absent count, etc.)')

    # Sheet 3
    doc.add_heading('Sheet 3: Alerts', level=2)
    doc.add_paragraph('Shows employees who need attention:')
    doc.add_paragraph('- High absences')
    doc.add_paragraph('- Frequent late arrivals')
    doc.add_paragraph('- Many missing punches')

    # Sheet 4
    doc.add_heading('Sheet 4: Penalties', level=2)
    doc.add_paragraph('This is your payroll deduction sheet with 25 columns:')

    penalties_table = doc.add_table(rows=26, cols=2)
    penalties_table.style = 'Table Grid'

    penalty_cols = [
        ('Column', 'Description'),
        ('A', 'CRM'),
        ('B', 'Name'),
        ('C', 'National ID'),
        ('D', 'Vendor'),
        ('E', 'PS ID'),
        ('F', 'Department'),
        ('G', 'Join Date'),
        ('H', 'Late Count'),
        ('I', 'Late Penalty (EGP)'),
        ('J', 'Missing Punches'),
        ('K', 'Punch Deduction (days)'),
        ('L', 'Absences'),
        ('M', 'Absence Deduction (days)'),
        ('N', 'Early Departure Count'),
        ('O', 'Early Departure Deduction'),
        ('P', 'Half Day Count'),
        ('Q', 'Half Day Deduction'),
        ('R', 'Sick Leave Count'),
        ('S', 'Sick Leave Deduction'),
        ('T', 'Unpaid Leave Count'),
        ('U', 'Unpaid Leave Deduction'),
        ('V', 'Total Penalty (EGP)'),
        ('W', 'Total Deduction (days)'),
        ('X', 'Warnings'),
        ('Y', 'Backdated Leaves'),
    ]

    for i, (col, desc) in enumerate(penalty_cols):
        row = penalties_table.rows[i]
        row.cells[0].text = col
        row.cells[1].text = desc
        if i == 0:
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ========== CHANGING STATUS ==========
    doc.add_heading('Changing Employee Status (Justifications)', level=1)
    doc.add_paragraph('You can change an employee\'s status directly in Excel!')

    doc.add_heading('How to Change Status', level=2)
    doc.add_paragraph('1. Open the generated Excel file')
    doc.add_paragraph('2. Go to Summary Report sheet')
    doc.add_paragraph('3. Click on any date cell (e.g., "Late" or "Absent")')
    doc.add_paragraph('4. A dropdown arrow appears')
    doc.add_paragraph('5. Click the arrow and select new status')

    p = doc.add_paragraph()
    p.add_run('Important: ').bold = True
    p.add_run('When you change a status in Summary Report, the Penalties sheet updates automatically!')

    # ========== DROPDOWN OPTIONS ==========
    doc.add_heading('Dropdown Status Options', level=1)

    # No Penalty
    doc.add_heading('No Penalty (Green)', level=2)
    no_penalty_table = doc.add_table(rows=11, cols=2)
    no_penalty_table.style = 'Table Grid'

    no_penalty = [
        ('Status', 'Use When'),
        ('Normal', 'Regular attendance, no issues'),
        ('Late (Approved)', 'Late arrival was approved by manager'),
        ('Early Departure (Approved)', 'Early leave was approved'),
        ('Annual Leave', 'Planned vacation leave'),
        ('Casual Leave', 'Personal/emergency leave'),
        ('Marriage Leave', 'Wedding leave'),
        ('Paternity Leave', 'Father\'s leave for new baby'),
        ('Maternity Leave', 'Mother\'s leave for new baby'),
        ('Bereavement Leave', 'Leave for family death'),
        ('Military Call Leave', 'Called for military service'),
    ]

    for i, (status, use) in enumerate(no_penalty):
        row = no_penalty_table.rows[i]
        row.cells[0].text = status
        row.cells[1].text = use
        if i == 0:
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True
        else:
            set_cell_shading(row.cells[0], '90EE90')

    doc.add_paragraph()

    # With Penalty
    doc.add_heading('With Penalty', level=2)
    with_penalty_table = doc.add_table(rows=10, cols=2)
    with_penalty_table.style = 'Table Grid'

    with_penalty = [
        ('Status', 'Penalty'),
        ('Late', '100/200/500 EGP (see scale below)'),
        ('Absent', '2 days deduction'),
        ('Missing Punch In', '0.5 day after 3 occurrences'),
        ('Missing Punch Out', '0.5 day after 3 occurrences'),
        ('Early Departure', '0.5 day deduction'),
        ('Half Day', '0.5 day deduction'),
        ('Early Leave (HD)', '0.5 day deduction'),
        ('Sick Leave', '0.25 day deduction'),
        ('Unpaid Leave', '1 day deduction'),
    ]

    penalty_colors = {
        'Late': 'FFFF00',
        'Absent': 'FF6B6B',
        'Missing Punch In': 'FFB6C1',
        'Missing Punch Out': 'FFB6C1',
        'Early Departure': 'FFA500',
        'Half Day': 'FFA500',
        'Early Leave (HD)': 'FFA500',
        'Sick Leave': 'ADD8E6',
        'Unpaid Leave': 'ADD8E6',
    }

    for i, (status, penalty) in enumerate(with_penalty):
        row = with_penalty_table.rows[i]
        row.cells[0].text = status
        row.cells[1].text = penalty
        if i == 0:
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True
        elif status in penalty_colors:
            set_cell_shading(row.cells[0], penalty_colors[status])

    doc.add_paragraph()

    # Backdated Leaves
    doc.add_heading('Backdated Leaves (Purple)', level=2)
    doc.add_paragraph('Use these for leaves from previous months:')

    bd_table = doc.add_table(rows=8, cols=2)
    bd_table.style = 'Table Grid'

    bd_leaves = [
        ('Status', 'Penalty'),
        ('Annual Leave (BD)', 'No deduction'),
        ('Casual Leave (BD)', 'No deduction'),
        ('Sick Leave (BD)', '0.25 day'),
        ('Unpaid Leave (BD)', '1 day'),
        ('Half Day (BD)', '0.5 day'),
        ('Early Departure (BD)', '0.5 day'),
        ('Early Leave (HD) (BD)', '0.5 day'),
    ]

    for i, (status, penalty) in enumerate(bd_leaves):
        row = bd_table.rows[i]
        row.cells[0].text = status
        row.cells[1].text = penalty
        if i == 0:
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True
        else:
            set_cell_shading(row.cells[0], 'E1D5E7')

    doc.add_paragraph()

    # ========== PENALTY RULES ==========
    doc.add_heading('Penalty Rules', level=1)

    doc.add_heading('Late Penalty Scale', level=2)
    late_table = doc.add_table(rows=5, cols=2)
    late_table.style = 'Table Grid'

    late_penalties = [
        ('Occurrence', 'Penalty'),
        ('1st Late', '100 EGP'),
        ('2nd Late', '200 EGP'),
        ('3rd Late', '500 EGP + Warning'),
        ('4th+ Late', '500 EGP each'),
    ]

    for i, (occ, pen) in enumerate(late_penalties):
        row = late_table.rows[i]
        row.cells[0].text = occ
        row.cells[1].text = pen
        if i == 0:
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True

    p = doc.add_paragraph()
    p.add_run('Note: ').bold = True
    p.add_run('"Late (Approved)" does NOT count toward late penalties!')

    doc.add_paragraph()

    doc.add_heading('Missing Punch Rules', level=2)
    punch_table = doc.add_table(rows=4, cols=2)
    punch_table.style = 'Table Grid'

    punch_rules = [
        ('Occurrence', 'Deduction'),
        ('1st - 3rd', 'No deduction'),
        ('4th - 5th', '0.5 day each'),
        ('6th+', '0.5 day + Warning'),
    ]

    for i, (occ, ded) in enumerate(punch_rules):
        row = punch_table.rows[i]
        row.cells[0].text = occ
        row.cells[1].text = ded
        if i == 0:
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True

    p = doc.add_paragraph()
    p.add_run('Note: ').bold = True
    p.add_run('All types count: Missing Punch In, Missing Punch In (Justified), Missing Punch Out, Missing Punch Out (Justified)')

    doc.add_paragraph()

    # ========== BACKDATED LEAVES ==========
    doc.add_heading('Adding Backdated Leaves', level=1)

    doc.add_heading('What are Backdated Leaves?', level=2)
    doc.add_paragraph('Leaves from a previous month that weren\'t recorded at the time.')

    p = doc.add_paragraph()
    p.add_run('Example: ').bold = True
    p.add_run('Employee took sick leave in November. You\'re processing January payroll (Dec 21 - Jan 20). You need to add this to the current payroll period.')

    doc.add_heading('How to Add Backdated Leave', level=2)

    p = doc.add_paragraph()
    p.add_run('Method 1: In the Leave Sheet (Before Report)').bold = True
    doc.add_paragraph('1. Open your Leave Sheet Excel file')
    doc.add_paragraph('2. Go to the current month\'s sheet (e.g., "Jan")')
    doc.add_paragraph('3. Find the employee\'s row')
    doc.add_paragraph('4. Select any date cell in the payroll period')
    doc.add_paragraph('5. Choose the (BD) version from dropdown (e.g., "Sick Leave (BD)")')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Method 2: In the Generated Report (After Report)').bold = True
    doc.add_paragraph('1. Open the generated Excel report')
    doc.add_paragraph('2. Go to Summary Report sheet')
    doc.add_paragraph('3. Click on any date cell for that employee')
    doc.add_paragraph('4. Select the (BD) version from dropdown')

    doc.add_heading('How to Identify Backdated Leaves', level=2)
    doc.add_paragraph('- Purple color in Summary Report')
    doc.add_paragraph('- Column Y in Penalties sheet shows count')
    doc.add_paragraph('- HR can review and verify these entries')

    # ========== TIPS ==========
    doc.add_heading('Tips & Tricks', level=1)

    p = doc.add_paragraph()
    p.add_run('Tip 1: Force Recalculate').bold = True
    doc.add_paragraph('If penalties don\'t update automatically, press Ctrl + Shift + F9 in Excel.')

    p = doc.add_paragraph()
    p.add_run('Tip 2: Remove Late Penalty').bold = True
    doc.add_paragraph('Change "Late" to "Late (Approved)" to remove the penalty.')

    p = doc.add_paragraph()
    p.add_run('Tip 3: Remove Early Departure Penalty').bold = True
    doc.add_paragraph('Change "Early Departure" to "Early Departure (Approved)".')

    p = doc.add_paragraph()
    p.add_run('Tip 4: Check Before Saving').bold = True
    doc.add_paragraph('Always check: Penalties sheet totals, Backdated Leaves column, Warning count.')

    p = doc.add_paragraph()
    p.add_run('Tip 5: Save Original').bold = True
    doc.add_paragraph('Keep a copy of the original generated report before making changes.')

    # ========== TROUBLESHOOTING ==========
    doc.add_heading('Troubleshooting', level=1)

    problems = [
        ("Application won't open", "Make sure you have the .exe file, not a shortcut. Try running as Administrator (right-click > Run as administrator)."),
        ('"File not found" error', "Make sure the Excel files are not open in another program. Close Excel and try again."),
        ("Filters not showing", "Load the Master Data file first. Filters appear after Master Data is loaded."),
        ("Leaves not appearing in report", "Make sure the CRM in Leave Sheet matches Master Data exactly. Check that the leave date is within the attendance date range."),
        ("Penalties not updating when I change status", "Press Ctrl + Shift + F9 to force recalculate. Make sure you're changing cells in the Summary Report sheet."),
        ("Report shows extra months", "This was fixed in v2.3. Make sure you're using the latest version."),
        ("Backdated leaves not counted in penalties", "Use the (BD) suffix versions. Check Column Y in Penalties sheet. Press Ctrl + Shift + F9 to recalculate."),
    ]

    for problem, solution in problems:
        p = doc.add_paragraph()
        p.add_run(f'Problem: {problem}').bold = True
        doc.add_paragraph(f'Solution: {solution}')
        doc.add_paragraph()

    # ========== QUICK REFERENCE ==========
    doc.add_page_break()
    doc.add_heading('Quick Reference Card', level=1)
    doc.add_paragraph('Print this page for your desk!')

    doc.add_heading('Step-by-Step', level=2)
    doc.add_paragraph('1. Open AttendanceDashboard.exe')
    doc.add_paragraph('2. Click "Select Attendance Files" - choose files')
    doc.add_paragraph('3. Click "Select Master Data" - choose file')
    doc.add_paragraph('4. Click "Select Leave Sheet" - choose file')
    doc.add_paragraph('5. Apply filters if needed')
    doc.add_paragraph('6. Click "Generate Report"')
    doc.add_paragraph('7. Save Excel file')

    doc.add_heading('To Change Status in Excel', level=2)
    doc.add_paragraph('1. Open generated Excel')
    doc.add_paragraph('2. Go to "Summary Report" sheet')
    doc.add_paragraph('3. Click date cell - click dropdown arrow')
    doc.add_paragraph('4. Select new status')
    doc.add_paragraph('5. Penalties update automatically!')

    doc.add_heading('Late Penalties', level=2)
    doc.add_paragraph('1st = 100 EGP | 2nd = 200 EGP | 3rd+ = 500 EGP')

    doc.add_heading('Deductions (Days)', level=2)
    doc.add_paragraph('Absent = 2 | Missing Punch (4th+) = 0.5 | Early Departure = 0.5')
    doc.add_paragraph('Half Day = 0.5 | Sick Leave = 0.25 | Unpaid Leave = 1')

    doc.add_heading('Keyboard Shortcut', level=2)
    p = doc.add_paragraph()
    p.add_run('Ctrl + Shift + F9').bold = True
    p.add_run(' = Force Recalculate in Excel')

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('User Guide - Attendance Dashboard v2.3')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.italic = True

    footer2 = doc.add_paragraph('Last Updated: January 2026')
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2.runs[0].font.italic = True

    # Save
    doc.save(r'C:\Users\high tech\Desktop\Attendance\V.5\files (1)\USER_GUIDE.docx')
    print("USER_GUIDE.docx created successfully!")

if __name__ == '__main__':
    create_user_guide()
